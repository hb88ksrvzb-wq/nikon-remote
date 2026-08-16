from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

doc = Document()

# ========== Page setup ==========
for section in doc.sections:
    section.orientation = WD_ORIENT.PORTRAIT
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.18)
    section.right_margin = Cm(3.18)

style = doc.styles['Normal']
font = style.font
font.name = '宋体'
font.size = Pt(12)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

def set_cell_border(cell, **kwargs):
    """Set cell border."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('start', 'top', 'end', 'bottom', 'insideH', 'insideV'):
        if edge in kwargs:
            element = OxmlElement(f'w:{edge}')
            for attr, val in kwargs[edge].items():
                element.set(qn(f'w:{attr}'), str(val))
            tcBorders.append(element)
    tcPr.append(tcBorders)

def set_paragraph_spacing(paragraph, before=0, after=0, line_spacing=1.5):
    """Set paragraph spacing."""
    pf = paragraph.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = line_spacing

def add_heading_styled(doc, text, level=1, font_name='黑体', font_size=None):
    """Add a styled heading."""
    heading = doc.add_heading(text, level=level)
    if level == 0:
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in heading.runs:
        run.font.name = font_name
        run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
        if font_size:
            run.font.size = Pt(font_size)
        run.font.color.rgb = RGBColor(0, 0, 0)
    return heading

def add_body_text(doc, text, bold=False, indent=False, font_name='宋体', font_size=12, align=None, color=None):
    """Add a body paragraph."""
    p = doc.add_paragraph()
    if align:
        p.alignment = align
    set_paragraph_spacing(p, before=0, after=3, line_spacing=1.5)
    if indent:
        p.paragraph_format.first_line_indent = Cm(0.74)
    run = p.add_run(text)
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    run.font.size = Pt(font_size)
    run.bold = bold
    if color:
        run.font.color.rgb = color
    return p

def add_image_centered(doc, img_path, width_inches=5.5, caption=""):
    """Add a centered image with optional caption."""
    if not os.path.exists(img_path):
        add_body_text(doc, f"[图片缺失: {img_path}]", font_size=10, color=RGBColor(128, 128, 128))
        return
    p_img = doc.add_paragraph()
    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p_img, before=6, after=3)
    run_img = p_img.add_run()
    run_img.add_picture(img_path, width=Inches(width_inches))
    
    if caption:
        p_cap = doc.add_paragraph()
        p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_paragraph_spacing(p_cap, before=0, after=6)
        run_cap = p_cap.add_run(caption)
        run_cap.font.name = '宋体'
        run_cap._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        run_cap.font.size = Pt(9)
        run_cap.font.color.rgb = RGBColor(100, 100, 100)
        run_cap.italic = True

def add_bullet(doc, text, level=0, bold_prefix=""):
    """Add a bullet point."""
    p = doc.add_paragraph(style='List Bullet')
    set_paragraph_spacing(p, before=0, after=2, line_spacing=1.5)
    p.paragraph_format.left_indent = Cm(0.74 * (level + 1))
    p.clear()
    if bold_prefix:
        run_b = p.add_run(bold_prefix)
        run_b.font.name = '宋体'
        run_b._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        run_b.font.size = Pt(12)
        run_b.bold = True
    run = p.add_run(text)
    run.font.name = '宋体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run.font.size = Pt(12)
    return p

# ========== COVER PAGE ==========
for _ in range(6):
    doc.add_paragraph()

title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run_t = title_p.add_run("基于大语言模型的智能会议纪要与任务分发系统")
run_t.font.name = '黑体'
run_t._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
run_t.font.size = Pt(26)
run_t.bold = True

doc.add_paragraph()

sub_p = doc.add_paragraph()
sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run_s = sub_p.add_run("用户操作手册")
run_s.font.name = '黑体'
run_s._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
run_s.font.size = Pt(22)

doc.add_paragraph()
doc.add_paragraph()

ver_p = doc.add_paragraph()
ver_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run_v = ver_p.add_run("版本号：V1.0")
run_v.font.name = '宋体'
run_v._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
run_v.font.size = Pt(14)

doc.add_paragraph()

dev_p = doc.add_paragraph()
dev_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run_d = dev_p.add_run("开发单位：北京科技创新有限公司")
run_d.font.name = '宋体'
run_d._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
run_d.font.size = Pt(14)

doc.add_paragraph()

date_p = doc.add_paragraph()
date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run_dt = date_p.add_run("编制日期：2026年6月")
run_dt.font.name = '宋体'
run_dt._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
run_dt.font.size = Pt(14)

doc.add_page_break()

# ========== TABLE OF CONTENTS ==========
add_heading_styled(doc, "目  录", level=0, font_name='黑体', font_size=18)
doc.add_paragraph()

toc_items = [
    ("一、", "引言", 1),
    ("", "1.1 编写目的", 2),
    ("", "1.2 项目背景", 2),
    ("", "1.3 术语与缩写解释", 2),
    ("二、", "软件概述", 1),
    ("", "2.1 软件名称与版本", 2),
    ("", "2.2 软件功能简介", 2),
    ("", "2.3 系统架构说明", 2),
    ("三、", "运行环境", 1),
    ("", "3.1 硬件环境要求", 2),
    ("", "3.2 软件环境要求", 2),
    ("", "3.3 网络环境要求", 2),
    ("四、", "软件的安装与卸载", 1),
    ("", "4.1 系统部署说明", 2),
    ("", "4.2 客户端访问说明", 2),
    ("五、", "软件操作指南", 1),
    ("", "5.1 用户登录", 2),
    ("", "5.2 会议音频上传与语音转写", 2),
    ("", "5.3 大模型智能会议摘要生成", 2),
    ("", "5.4 会议待办事项提取与任务分发", 2),
    ("", "5.5 历史会议记录检索与管理", 2),
    ("", "5.6 系统参数与模型配置", 2),
    ("六、", "常见问题与故障处理", 1),
    ("七、", "技术支持与联系方式", 1),
]

for prefix, item, level in toc_items:
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=0, after=2, line_spacing=1.8)
    indent_size = Cm(0.74 * (level - 1))
    p.paragraph_format.left_indent = indent_size
    run = p.add_run(f"{prefix}{item}")
    run.font.name = '宋体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run.font.size = Pt(14) if level == 1 else Pt(12)
    run.bold = (level == 1)

doc.add_page_break()

# ========== CHAPTER 1: INTRODUCTION ==========
add_heading_styled(doc, "一、引言", level=1)

add_heading_styled(doc, "1.1 编写目的", level=2)
add_body_text(doc, "本文档是《基于大语言模型的智能会议纪要与任务分发系统 V1.0》（以下简称“本系统”）的用户操作手册，旨在为用户提供全面、详细的软件使用指导。手册内容涵盖系统的功能特点、运行环境、安装部署流程以及各功能模块的详细操作步骤，帮助用户快速掌握本系统的使用方法，充分发挥系统在会议纪要自动生成与任务智能分发方面的核心价值。", indent=True)

add_heading_styled(doc, "1.2 项目背景", level=2)
add_body_text(doc, "随着企业信息化建设的不断深入，日常工作中产生的会议数量日益增多。传统的会议纪要整理方式依赖人工记录与手动汇总，存在效率低下、信息遗漏、任务落实不清晰等痛点。", indent=True)
add_body_text(doc, "本系统基于先进的大语言模型（Large Language Model, LLM）技术，结合高精度自动语音识别（ASR）引擎，实现了从会议音频上传、语音转文字、智能摘要生成到待办任务自动提取与分发的一站式闭环解决方案。系统能够自动将冗长的会议录音转化为结构化的会议纪要和可追踪的待办事项清单，大幅提升团队协作效率与会议产出质量。", indent=True)

add_heading_styled(doc, "1.3 术语与缩写解释", level=2)

table = doc.add_table(rows=6, cols=3, style='Table Grid')
table.autofit = True
headers = ['术语/缩写', '全称', '解释说明']
for i, h in enumerate(headers):
    cell = table.rows[0].cells[i]
    cell.text = h
    for p in cell.paragraphs:
        for run in p.runs:
            run.font.bold = True
            run.font.size = Pt(10)
            run.font.name = '宋体'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

data = [
    ['LLM', 'Large Language Model', '大语言模型，用于理解自然语言并生成会议摘要及任务'],
    ['ASR', 'Automatic Speech Recognition', '自动语音识别技术，将会议音频转化为文本'],
    ['NLP', 'Natural Language Processing', '自然语言处理，用于分析文本语义、提取关键信息'],
    ['API', 'Application Programming Interface', '应用程序编程接口，系统各模块间通信的标准协议'],
    ['SDK', 'Software Development Kit', '软件开发工具包，供第三方系统集成的开发组件'],
]
for r_idx, row_data in enumerate(data):
    for c_idx, val in enumerate(row_data):
        cell = table.rows[r_idx + 1].cells[c_idx]
        cell.text = val
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.size = Pt(10)
                run.font.name = '宋体'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

doc.add_page_break()

# ========== CHAPTER 2: SOFTWARE OVERVIEW ==========
add_heading_styled(doc, "二、软件概述", level=1)

add_heading_styled(doc, "2.1 软件名称与版本", level=2)
add_body_text(doc, "软件全称：基于大语言模型的智能会议纪要与任务分发系统")
add_body_text(doc, "软件简称：智能会议助手")
add_body_text(doc, "当前版本：V1.0")
add_body_text(doc, "软件类型：Web应用系统（B/S架构）")
add_body_text(doc, "开发语言：Python 3.9+ / TypeScript / React 18")
add_body_text(doc, "数据库：PostgreSQL 15 + Redis 7.0")

add_heading_styled(doc, "2.2 软件功能简介", level=2)
add_body_text(doc, "本系统围绕“会议前-会议中-会议后”的全流程管理理念，提供以下核心功能模块：", indent=True)

funcs = [
    ("用户认证与权限管理：", "支持多角色（管理员、普通用户、访客）的注册、登录与权限控制，保障系统数据安全。"),
    ("音频上传与实时语音转写：", "支持多种音频格式（mp3、wav、m4a等）上传，基于高精度ASR引擎实现实时语音识别，自动标注说话人标签与时间戳。"),
    ("大模型智能摘要生成：", "利用先进的大语言模型对会议转录文本进行深度语义理解，自动提炼会议核心议题、关键讨论要点、主要结论与决议事项，生成结构化会议纪要。"),
    ("待办任务智能提取与分发：", "自动从会议纪要中识别待办事项（Action Items），智能匹配负责人，设定优先级与截止日期，并支持一键同步至企业协作平台（钉钉、飞书等）。"),
    ("历史会议记录管理：", "提供按时间、主题、参会人等维度的会议记录检索功能，支持全文关键词搜索，方便用户随时回顾历史会议内容。"),
    ("提示词模板与模型配置：", "支持管理员自定义摘要生成的提示词模板（Prompt Template），灵活调整不同会议场景下的生成策略，适配研发、市场、行政等不同部门的会议风格。"),
]

for prefix, detail in funcs:
    add_bullet(doc, detail, bold_prefix=prefix)

add_heading_styled(doc, "2.3 系统架构说明", level=2)
add_body_text(doc, "本系统采用前后端分离的B/S（Browser/Server）架构，整体分为以下四层：", indent=True)

add_bullet(doc, "浏览器进行交互操作，通过HTTPS协议与后端API进行数据通信。", bold_prefix="表现层（前端）：", level=0)
add_bullet(doc, "处理业务逻辑，包括用户认证、音频管理、ASR调度、LLM调用及任务分发等模块，采用RESTful API设计。", bold_prefix="业务逻辑层（后端）：", level=0)
add_bullet(doc, "负责数据的持久化存储，使用PostgreSQL管理结构化数据（用户、会议、任务），使用Redis缓存高频访问数据。", bold_prefix="数据存储层：", level=0)
add_bullet(doc, "集成第三方API，主要包括ASR语音转写服务和大语言模型（如GPT-4o等）的API调用。", bold_prefix="服务集成层：", level=0)

doc.add_page_break()

# ========== CHAPTER 3: RUNNING ENVIRONMENT ==========
add_heading_styled(doc, "三、运行环境", level=1)

add_heading_styled(doc, "3.1 硬件环境要求（服务器端）", level=2)
table_hw = doc.add_table(rows=6, cols=3, style='Table Grid')
table_hw.autofit = True
hw_headers = ['硬件项目', '最低配置', '推荐配置']
for i, h in enumerate(hw_headers):
    cell = table_hw.rows[0].cells[i]
    cell.text = h
    for p in cell.paragraphs:
        for run in p.runs:
            run.font.bold = True
            run.font.size = Pt(10)
            run.font.name = '宋体'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

hw_data = [
    ['CPU', '4核 2.0GHz', '8核 3.0GHz 及以上'],
    ['内存', '8GB', '16GB 及以上'],
    ['硬盘', '50GB 可用空间', '200GB SSD 及以上'],
    ['GPU（可选）', '无（使用云端API）', 'NVIDIA A10/A100（本地LLM部署）'],
    ['网络', '10Mbps 宽带', '100Mbps 光纤'],
]
for r_idx, row_data in enumerate(hw_data):
    for c_idx, val in enumerate(row_data):
        cell = table_hw.rows[r_idx + 1].cells[c_idx]
        cell.text = val
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.size = Pt(10)
                run.font.name = '宋体'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

add_heading_styled(doc, "3.2 软件环境要求", level=2)

table_sw = doc.add_table(rows=8, cols=3, style='Table Grid')
table_sw.autofit = True
sw_headers = ['软件类别', '软件名称', '版本要求']
for i, h in enumerate(sw_headers):
    cell = table_sw.rows[0].cells[i]
    cell.text = h
    for p in cell.paragraphs:
        for run in p.runs:
            run.font.bold = True
            run.font.size = Pt(10)
            run.font.name = '宋体'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

sw_data = [
    ['操作系统（服务器）', 'Linux (Ubuntu/CentOS)', 'Ubuntu 22.04 LTS 或 CentOS 8+'],
    ['Web服务器', 'Nginx', '1.24+'],
    ['应用服务器', 'Gunicorn + Uvicorn', 'Gunicorn 21.2+, Uvicorn 0.24+'],
    ['数据库', 'PostgreSQL', '15.x'],
    ['缓存数据库', 'Redis', '7.0+'],
    ['客户端浏览器', 'Chrome / Edge / Firefox', 'Chrome 120+, Edge 120+, Firefox 121+'],
    ['Python运行时', 'Python', '3.9+'],
]
for r_idx, row_data in enumerate(sw_data):
    for c_idx, val in enumerate(row_data):
        cell = table_sw.rows[r_idx + 1].cells[c_idx]
        cell.text = val
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.size = Pt(10)
                run.font.name = '宋体'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

add_heading_styled(doc, "3.3 网络环境要求", level=2)
add_bullet(doc, "服务器需具备公网IP或通过内网DNS解析，确保客户端可访问。")
add_bullet(doc, "需向外部API服务（ASR引擎、LLM服务商等）发起HTTPS请求，因此服务器需开放443端口的出站访问。")
add_bullet(doc, "推荐使用HTTPS协议进行数据传输，建议部署SSL/TLS证书以保障通信安全。")
add_bullet(doc, "本系统首页访问地址示例：https://meeting-assistant.ai")

doc.add_page_break()

# ========== CHAPTER 4: INSTALLATION ==========
add_heading_styled(doc, "四、软件的安装与卸载", level=1)

add_heading_styled(doc, "4.1 系统部署说明", level=2)
add_body_text(doc, "本系统为B/S架构的Web应用，服务器端部署步骤如下：", indent=True)

steps = [
    "步骤一：确保服务器满足第三章所述的软硬件环境要求，安装Python 3.9+、PostgreSQL 15、Redis 7.0。",
    "步骤二：将项目源码部署包解压至服务器目标目录（如 /opt/meeting_assistant/）。",
    "步骤三：进入项目根目录，执行 pip install -r requirements.txt 安装Python依赖包。",
    "步骤四：编辑配置文件 config.yaml，配置数据库连接地址、ASR服务API密钥、大语言模型API密钥等参数。",
    "步骤五：执行数据库初始化脚本 python manage.py init_db，自动创建所需的数据表结构。",
    "步骤六：使用 Gunicorn + Uvicorn 启动后端服务，并使用 Nginx 进行反向代理配置。",
    "步骤七：访问系统首页，使用默认管理员账号（admin / admin123456）登录，验证系统是否正常运行。",
    "步骤八：登录后请立即修改默认管理员密码，确保系统安全。",
]
for s in steps:
    add_body_text(doc, s, indent=True)

add_heading_styled(doc, "4.2 客户端访问说明", level=2)
add_body_text(doc, "本系统为纯Web应用，客户端无需安装任何专用软件，仅需通过主流浏览器（Chrome、Edge、Firefox等）访问系统URL即可。首次登录后，建议用户根据角色分配修改个人密码和基本信息。", indent=True)
add_body_text(doc, "若需卸载本系统，仅需停止服务器端进程（Gunicorn），删除部署目录及相关数据库即可完成清理。", indent=True)

doc.add_page_break()

# ========== CHAPTER 5: OPERATION GUIDE ==========
add_heading_styled(doc, "五、软件操作指南", level=1)

add_body_text(doc, "本章详细说明本系统各功能模块的操作方法，并配以界面截图辅助说明。", indent=True)

# 5.1 Login
add_heading_styled(doc, "5.1 用户登录", level=2)
add_body_text(doc, "用户通过浏览器访问系统首页后，进入登录界面。在登录框中输入已分配的用户名/邮箱和密码，点击“登录”按钮即可进入系统主工作台。", indent=True)
add_body_text(doc, "系统支持以下登录方式：", indent=True)
add_bullet(doc, "用户名/邮箱 + 密码登录（默认方式）")
add_bullet(doc, "企业SSO单点登录（需管理员预先配置LDAP/OAuth2.0对接）")
add_bullet(doc, "手机验证码快捷登录（需管理员开启短信服务配置）")

add_image_centered(doc, "img/login_page.png", width_inches=4.5, caption="图5.1-1 用户登录界面")

# 5.2 Upload
add_heading_styled(doc, "5.2 会议音频上传与语音转写", level=2)
add_body_text(doc, "登录系统后，左侧导航栏默认选中“🎙️ 语音上传转写”模块。用户可通过以下方式上传会议音频：", indent=True)
add_bullet(doc, "拖拽上传：将本地音频文件（mp3、wav、m4a等）直接拖入页面中的上传区域。")
add_bullet(doc, "点击上传：点击上传区域，弹出文件选择对话框，选择本地音频文件。")
add_bullet(doc, "批量上传：支持一次性选择多个会议音频文件，系统将对每个文件分别建立转写任务。")
add_body_text(doc, "文件上传完成后，系统自动启动语音转写任务。上传区域下方会显示转写进度条与状态信息（如“正在转写中”、“已完成”、“转写失败”等）。页面右侧面板将实时展示语音识别生成的文本预览，包含说话人标签和时间戳信息。", indent=True)
add_body_text(doc, "转写配置说明：", indent=True, bold=True)
add_bullet(doc, "说话人识别：系统默认开启，支持2-5人的说话人自动标注（Diarization）。")
add_bullet(doc, "ASR引擎选择：默认使用高精双工混响引擎V2，支持中英文混合识别。")
add_bullet(doc, "热词字典：系统预置了互联网技术、软件研发等行业热词字典，用户也可上传自定义热词表以提升特定领域的识别准确率。")

add_image_centered(doc, "img/dashboard_upload.png", width_inches=5.0, caption="图5.2-1 音频上传与实时转写界面")

# 5.3 Summary
add_heading_styled(doc, "5.3 大模型智能会议摘要生成", level=2)
add_body_text(doc, "当语音转写任务完成后，用户可点击导航栏中的“📝 大模型智能摘要”模块进入摘要生成页面。", indent=True)
add_body_text(doc, "操作步骤：", indent=True, bold=True)
add_bullet(doc, "选择模型：在页面顶部工具栏中选择当前使用的大语言模型（如GPT-4o精调版、文心一言等）。")
add_bullet(doc, "选择提示词模板：选择适合当前会议类型的摘要生成提示词模板（如“研发团队详细会议纪要”、“市场部日常周会”等）。")
add_bullet(doc, "生成摘要：点击“重新生成”按钮，系统将对应会议的语音转录文本送入大语言模型，自动生成结构化的会议纪要。")
add_bullet(doc, "结果确认：左侧面板显示原始转录文本供用户参考比对；右侧面板展示智能生成的会议纪要，包括会议名称、时间、核心议程与讨论要点、关键结论与决议等内容。")
add_bullet(doc, "一键同步：确认纪要无误后，点击“确认并一键同步到待办任务”按钮，系统将自动提取会议中的待办事项并跳转至任务分发模块。")

add_image_centered(doc, "img/ai_summary.png", width_inches=5.0, caption="图5.3-1 大模型智能摘要生成界面")

# 5.4 Task
add_heading_styled(doc, "5.4 会议待办事项提取与任务分发", level=2)
add_body_text(doc, "通过大模型摘要模块自动提取的待办事项，或用户手动添加的任务，均集中在“📋 待办任务分发”模块中进行统一管理。", indent=True)
add_body_text(doc, "功能要点：", indent=True, bold=True)
add_bullet(doc, "自动提取：AI算法自动从会议纪要文本中识别出待办事项（Action Items），并智能匹配候选负责人。")
add_bullet(doc, "手动调整：支持手动修改任务标题、指派人、优先级（紧急/高、普通/中、低优先）、截止日期等属性。")
add_bullet(doc, "优先级颜色标识：紧急/高优先级任务以红色标识，普通/中优先级以橙色标识，低优先级以灰色标识，一目了然。")
add_bullet(doc, "手动添加：点击“手动添加待办任务”按钮，可录入会议纪要中未自动识别的额外任务。")
add_bullet(doc, "一键同步：点击“一键同步至企业钉钉/飞书待办”按钮，系统将通过API接口将任务列表自动推送至企业协作平台，确保任务落实到位。")

add_image_centered(doc, "img/task_distribution.png", width_inches=5.0, caption="图5.4-1 待办任务提取与分发界面")

# 5.5 History
add_heading_styled(doc, "5.5 历史会议记录检索与管理", level=2)
add_body_text(doc, "点击导航栏中的“📂 历史会议记录”模块，进入历史会议纪要检索与管理中心。", indent=True)
add_body_text(doc, "主要功能：", indent=True, bold=True)
add_bullet(doc, "智能检索：支持输入会议名称、议程、关键词进行全文模糊搜索，快速定位目标会议记录。")
add_bullet(doc, "日期筛选：支持按会议召开日期范围进行过滤。")
add_bullet(doc, "列表展示：以表格形式展示所有历史会议记录，包括会议主题、召开时间、会议时长、提取的待办事项数量及状态。")
add_bullet(doc, "操作管理：每条记录均提供“查看”、“导出”、“删除”操作按钮。点击“查看”可跳转至该次会议的完整纪要详情页；点击“导出”可下载Word或PDF格式的会议纪要文档；点击“删除”可移除不需要的会议记录。")
add_bullet(doc, "分页浏览：支持分页功能，默认每页显示10条记录，可通过翻页按钮浏览全部历史会议。")

add_image_centered(doc, "img/history_records.png", width_inches=5.0, caption="图5.5-1 历史会议记录检索与管理界面")

# 5.6 Settings
add_heading_styled(doc, "5.6 系统参数与模型配置", level=2)
add_body_text(doc, "系统管理员可通过导航栏中的“⚙️ 系统参数设置”模块进行全局配置管理，主要包括：", indent=True)
add_bullet(doc, "大模型API配置：配置对接的大语言模型服务商（如OpenAI、百度文心、阿里通义等）的API地址与密钥。")
add_bullet(doc, "提示词模板管理：新增、编辑、删除不同场景下的会议摘要生成提示词模板（Prompt Template），支持变量占位符（如{meeting_type}、{participants}等）。")
add_bullet(doc, "ASR引擎参数：调整语音识别的语言模型、声学模型、热词字典等参数。")
add_bullet(doc, "用户与权限管理：管理系统用户账号、角色分配（管理员/普通用户/访客）及相应的操作权限。")
add_bullet(doc, "企业协作平台对接：配置钉钉、飞书、企业微信等第三方平台的API集成参数。")
add_bullet(doc, "日志与审计：查看系统操作日志，追踪关键操作记录，满足安全审计需求。")

doc.add_page_break()

# ========== CHAPTER 6: FAQ ==========
add_heading_styled(doc, "六、常见问题与故障处理", level=1)

faqs = [
    ("Q1：音频上传后长时间没有转写结果？", "A：请检查网络连接是否正常，确认服务器与ASR服务的通信畅通。若网络正常，请检查音频文件格式是否被支持（支持mp3、wav、m4a等），文件大小是否超过500MB限制。如问题仍存在，请联系系统管理员查看后台转写任务队列。"),
    ("Q2：大模型生成的摘要内容不准确或遗漏关键信息？", "A：可尝试调整提示词模板（Prompt Template），针对不同类型的会议使用更合适的提示词。也可以尝试切换不同的大语言模型进行生成。若持续存在质量问题，请在系统设置中提交反馈日志，管理员将优化默认提示词策略。"),
    ("Q3：任务未能成功同步至钉钉/飞书？", "A：首先确认系统设置中已正确配置企业协作平台的API集成参数（AppKey、AppSecret等）。其次检查相关账号是否具有在企业平台中创建任务的权限。如配置无误，请检查系统操作日志中的API调用错误信息。"),
    ("Q4：浏览器页面加载缓慢或白屏？", "A：建议清除浏览器缓存后重试，推荐使用Chrome 120+或Edge 120+版本浏览器。若问题持续，请检查网络延迟和服务器负载状态。"),
    ("Q5：忘记登录密码怎么办？", "A：在登录页面点击“忘记密码”链接，按照提示通过注册邮箱或手机号重置密码。若无法自助重置，请联系系统管理员进行手动密码重置。"),
]

for q, a in faqs:
    add_body_text(doc, q, bold=True)
    add_body_text(doc, a, indent=True)
    doc.add_paragraph()

doc.add_page_break()

# ========== CHAPTER 7: SUPPORT ==========
add_heading_styled(doc, "七、技术支持与联系方式", level=1)

add_body_text(doc, "如您在使用本系统过程中遇到任何问题，或希望获取更多技术支持和定制化服务，请通过以下方式联系我们：", indent=True)
doc.add_paragraph()

contact_items = [
    ("技术支持热线：", "400-XXX-XXXX（工作日 9:00 - 18:00）"),
    ("技术支持邮箱：", "support@tech-innovation.com"),
    ("官方网站：", "https://www.tech-innovation.com"),
    ("在线帮助文档：", "https://docs.meeting-assistant.ai"),
    ("企业微信服务号：", "北京科技创新有限公司"),
    ("联系地址：", "北京市海淀区中关村科技园区创新大厦A座18层"),
]

for prefix, detail in contact_items:
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=0, after=3, line_spacing=1.5)
    run_prefix = p.add_run(prefix)
    run_prefix.font.name = '宋体'
    run_prefix._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run_prefix.font.size = Pt(12)
    run_prefix.bold = True
    run_detail = p.add_run(detail)
    run_detail.font.name = '宋体'
    run_detail._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run_detail.font.size = Pt(12)

doc.add_paragraph()
doc.add_paragraph()

end_p = doc.add_paragraph()
end_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run_end = end_p.add_run("— 文档结束 —")
run_end.font.name = '宋体'
run_end._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
run_end.font.size = Pt(12)
run_end.font.color.rgb = RGBColor(128, 128, 128)

# ========== Save ==========
output_path = "基于大语言模型的智能会议纪要与任务分发系统V1.0_用户操作手册.docx"
doc.save(output_path)
print(f"Word document generated successfully: {output_path}")
